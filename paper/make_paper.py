"""Generate the paper skeleton as a .docx.

Methods are written out in full; results/abstract carry clearly labelled
[PLACEHOLDER] markers where real numbers and figures go. The structure follows
what a venue like JCTC / JCIM / Bioinformatics expects for a methods paper whose
central claim is an ablation.
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches

PLACEHOLDER = RGBColor(0xB0, 0x00, 0x00)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def ph(paragraph, text):
    """Add a red [PLACEHOLDER] run."""
    run = paragraph.add_run(text)
    run.font.color.rgb = PLACEHOLDER
    run.bold = True
    return run


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def build():
    doc = Document()

    # Base styles
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Ramachandran-regularized torsion-space diffusion improves the physical "
        "fidelity of generated intrinsically disordered protein ensembles"
    )
    r.bold = True
    r.font.size = Pt(15)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph(sub, "[AUTHORS] — [AFFILIATIONS] — [CORRESPONDING EMAIL]")

    # ---- Abstract ----
    add_heading(doc, "Abstract", 1)
    p = doc.add_paragraph()
    p.add_run(
        "Intrinsically disordered proteins (IDPs) populate broad conformational "
        "ensembles that are central to their function yet difficult to model. "
        "Generative models in backbone torsion space are attractive because they "
        "sidestep the rigid-body constraints of Cartesian generation, but they can "
        "place mass in physically implausible regions of (φ, ψ) space. We introduce "
        "a torus-native wrapped-diffusion model for IDP torsion ensembles, "
        "conditioned on sequence through ESM-2 embeddings, and add a "
        "residue-type-specific Ramachandran kernel-density regularizer that pulls "
        "generated angles toward empirically populated basins. We evaluate against "
        "a residue-specific statistical-coil null model and idpGAN on held-out "
        "Protein Ensemble Database (PED) proteins, scoring per-residue (φ, ψ) "
        "Jensen–Shannon divergence, distributions of global size observables, and "
        "agreement with experimental SAXS and NMR data. "
    )
    ph(p, "[RESULT: headline effect of the regularizer on JSD and SAXS/NMR χ², with "
          "the fidelity/diversity trade-off quantified — fill from runs/*/eval.json.]")

    # ---- 1. Introduction ----
    add_heading(doc, "1. Introduction", 1)
    body(doc,
         "IDPs and intrinsically disordered regions are abundant in eukaryotic "
         "proteomes and mediate signalling, regulation, and condensate formation. "
         "Because they lack a single native fold, their behaviour is described by "
         "an ensemble of interconverting conformations rather than one structure. "
         "Experimentally, such ensembles are constrained by small-angle X-ray "
         "scattering (SAXS), nuclear magnetic resonance (NMR) chemical shifts, "
         "residual dipolar couplings, and paramagnetic relaxation enhancements, and "
         "are deposited in resources such as the Protein Ensemble Database (PED).")
    body(doc,
         "Generative modelling of IDP ensembles has the potential to accelerate "
         "this characterisation. Models that operate directly on backbone torsion "
         "angles (φ, ψ) are appealing because local bond geometry is approximately "
         "fixed, so torsions are the natural degrees of freedom; however, angles are "
         "periodic, and naively treating them as Euclidean quantities — for example "
         "regressing sin/cos features under a mean-squared error — introduces "
         "artefacts at the ±π seam and can populate disallowed Ramachandran "
         "regions. A second difficulty is that a model trained only to reproduce "
         "marginal angle statistics may still generate sterically or "
         "stereochemically implausible conformations.")
    body(doc,
         "We make the following contributions. (i) We formulate IDP torsion "
         "generation as torus-native wrapped diffusion with an x0 (clean-angle) "
         "parameterisation, which respects angular periodicity by construction. "
         "(ii) We introduce a residue-type-specific Ramachandran kernel-density "
         "regularizer and study its effect as the paper's central question. "
         "(iii) We evaluate with protein-level data splits against both a "
         "statistical-coil null and idpGAN, and validate against experimental SAXS "
         "and NMR observables, reporting fidelity and diversity jointly.")

    # ---- 2. Related work ----
    add_heading(doc, "2. Related work", 1)
    body(doc,
         "Statistical-coil and fragment-assembly models (e.g. flexible-meccano-style "
         "approaches) generate disordered ensembles by sampling residue-specific "
         "backbone angles independently and remain a standard null model. Learned "
         "generative approaches include idpGAN, a sequence-conditioned generator of "
         "coarse-grained IDP conformations, and a growing set of diffusion models for "
         "protein structure. Torsional diffusion introduced diffusion on the torus "
         "for small-molecule conformer generation, establishing the wrapped-diffusion "
         "machinery we adapt here to protein backbones.")
    p = doc.add_paragraph()
    ph(p, "[Expand with the specific prior-work citations you will include; "
          "position this work relative to each — one or two sentences each.]")

    # ---- 3. Methods ----
    add_heading(doc, "3. Methods", 1)

    add_heading(doc, "3.1 Torsion representation and circular geometry", 2)
    body(doc,
         "Each residue i contributes a backbone dihedral pair (φ_i, ψ_i) on the "
         "circle S^1, so a chain of L residues is a point on the 2L-torus. Angles "
         "are stored in (−π, π]. All operations that compare angles use the signed "
         "shortest-arc difference Δ(a, b) = wrap(a − b), and the network consumes "
         "the periodicity-respecting feature encoding "
         "[cos φ, sin φ, cos ψ, sin ψ]. Terminal residues, whose φ or ψ is "
         "undefined, are masked throughout training and evaluation.")

    add_heading(doc, "3.2 Torus-native wrapped diffusion", 2)
    body(doc,
         "The forward process adds wrapped-normal noise of increasing width to the "
         "clean angles: q(x_t | x_0) wraps a Gaussian of standard deviation σ_t onto "
         "the circle, where σ_t follows a cosine schedule rescaled so that the "
         "highest-noise level is effectively uniform on the circle. As σ_t grows, "
         "the wrapped normal converges to the uniform distribution, which is the "
         "prior from which the reverse process is initialised.")
    body(doc,
         "We use an x0-prediction parameterisation: the denoiser predicts the clean "
         "angles rather than the added noise. On the circle the winding number of "
         "the noise is unidentifiable after wrapping, so noise prediction is "
         "ill-posed; predicting the destination angle avoids this and makes the "
         "Ramachandran regularizer straightforward to apply, because it acts on a "
         "concrete predicted angle. The reverse step interpolates along the "
         "shortest arc from x_t toward the predicted clean angle by a fraction set "
         "by the drop in angular variance between consecutive steps, then adds "
         "wrapped posterior noise; every operation is expressed through the "
         "shortest-arc difference so nothing crosses the ±π seam incorrectly.")

    add_heading(doc, "3.3 Sequence-conditioned denoiser", 2)
    body(doc,
         "The denoiser is a pre-norm Transformer encoder. Per-residue inputs are a "
         "linear projection of the noised (cos, sin) features added to a projection "
         "of the per-residue ESM-2 embedding; conditioning is injected per residue "
         "rather than pooled, because torsion angles are local quantities. The "
         "diffusion timestep is embedded sinusoidally and supplied through adaptive "
         "layer-norm modulation in every block. Padding is handled with a "
         "key-padding mask so variable-length chains batch correctly. The default "
         "model has ")
    p = doc.paragraphs[-1]
    ph(p, "[N] ")
    p.add_run(
        "parameters (width 384, depth 6), sized for a single consumer GPU; "
        "width and depth are configuration options for scaling up.")

    add_heading(doc, "3.4 Ramachandran-KDE regularizer (central contribution)", 2)
    body(doc,
         "We precompute a kernel-density estimate of the empirical (φ, ψ) "
         "distribution, separately for each residue class — general, glycine, "
         "proline, and pre-proline — because these classes occupy distinct basins. "
         "Densities are represented on a periodic grid over (−π, π]^2 built from von "
         "Mises kernels, and are queried with wrap-around bilinear interpolation so "
         "the density is differentiable in the query angle. Crucially, the KDE is "
         "estimated from disordered-protein data (PED), not from folded-protein "
         "Ramachandran maps, so it reflects the PPII-enriched IDP landscape rather "
         "than biasing the model toward folded statistics.")
    body(doc,
         "During training we add to the angular reconstruction loss a penalty equal "
         "to the mean negative log-density, under the appropriate per-residue class "
         "density, of the model's predicted clean angles at unmasked residues. The "
         "penalty weight λ is annealed from λ_max to λ_min on a cosine schedule: "
         "early in training the prior stabilises learning and shepherds predictions "
         "toward populated regions, while late in training the data dominate so the "
         "model can learn sequence-specific deviations and the prior cannot, by "
         "itself, collapse ensemble diversity. The 'without-regularization' arm sets "
         "λ_max = 0; all else is identical.")
    body(doc,
         "The training objective is the masked angular reconstruction loss "
         "L_recon = mean over residues of [1 − cos Δφ] + [1 − cos Δψ], plus the "
         "annealed regularizer λ(t) · L_KDE. The 1 − cos loss is the natural smooth, "
         "periodic discrepancy on the circle and is minimised when predicted and "
         "true angles coincide.")

    add_heading(doc, "3.5 Reconstruction and observables", 2)
    body(doc,
         "To compute global observables and to interface with experimental forward "
         "models, we reconstruct N, CA, C backbone coordinates from generated "
         "torsions using the Natural Extension Reference Frame algorithm with ideal "
         "bond lengths and angles and trans-ω. From the CA trace we compute the "
         "radius of gyration, end-to-end distance, and asphericity. Because "
         "torsion-only reconstruction with ideal geometry cannot detect steric "
         "clashes, we additionally report a CA–CA clash rate to quantify this known "
         "limitation rather than leaving it implicit.")

    add_heading(doc, "3.6 Data, splits, and baselines", 2)
    body(doc,
         "We use PED ensembles, treating each entry as one protein whose deposited "
         "models constitute its reference ensemble. Sequence embeddings are taken "
         "from ESM-2 (650M) and cached per sequence. Data are split at the level of "
         "whole proteins — optionally grouped by sequence-identity clusters — so that "
         "no conformer of a test protein is seen in training; a per-conformer split "
         "would leak ensemble members across the boundary. We compare against a "
         "residue-specific statistical-coil null model that samples each residue's "
         "angles independently from the same per-class densities, and against "
         "idpGAN on the same held-out proteins.")
    p = doc.add_paragraph()
    ph(p, "[DATA: number of PED entries, total conformers, length distribution, "
          "clustering threshold, and train/val/test counts — fill from prepare_data.py output.]")

    add_heading(doc, "3.7 Evaluation metrics", 2)
    body(doc,
         "We report (i) per-residue (φ, ψ) Jensen–Shannon divergence between "
         "generated and reference ensembles on a shared 2-D angle grid; (ii) "
         "1-Wasserstein distances between generated and reference distributions of "
         "radius of gyration, end-to-end distance, and asphericity, since IDP "
         "behaviour is defined by ensemble spread rather than means; (iii) ensemble "
         "diversity as the mean pairwise circular distance, reported alongside "
         "fidelity so the regularizer's trade-off is explicit; and (iv) reduced "
         "χ² against experimental SAXS profiles and NMR chemical shifts, computed "
         "from ensemble-averaged forward predictions on held-out proteins.")

    # ---- 4. Results ----
    add_heading(doc, "4. Results", 1)

    add_heading(doc, "4.1 The regularizer's effect on ensemble fidelity (central result)", 2)
    p = doc.add_paragraph()
    ph(p, "[RESULT: with vs without regularization on held-out proteins. Report "
          "per-residue JSD (mean/median), observable Wasserstein distances, and the "
          "fidelity/diversity trade-off. State plainly whether the regularizer helps, "
          "by how much, and what diversity is sacrificed. Source: runs/with_reg/eval.json "
          "vs runs/without_reg/eval.json.]")
    _results_table(doc,
        "Table 1. Held-out fidelity, with vs without Ramachandran regularization.",
        ["Metric", "Without reg", "With reg", "Stat. coil"],
        ["Per-residue JSD (↓)", "Rg Wasserstein, Å (↓)", "Ree Wasserstein, Å (↓)",
         "Asphericity Wasserstein (↓)", "Ensemble diversity", "CA–CA clash rate (↓)"])

    add_heading(doc, "4.2 Comparison to baselines", 2)
    p = doc.add_paragraph()
    ph(p, "[RESULT: model vs statistical-coil null vs idpGAN on identical held-out "
          "proteins and identical metrics. The key sentence to support: the learned "
          "model beats the independent-residue null on ensemble observables, not just "
          "marginals. See docs/baselines.md for the idpGAN conversion path.]")

    add_heading(doc, "4.3 Agreement with experimental observables", 2)
    p = doc.add_paragraph()
    ph(p, "[RESULT: reduced χ² vs SAXS (Rg and/or full profile) and ≥1 NMR observable "
          "on held-out proteins, with and without regularization. This is the most "
          "load-bearing result for acceptance. Source: idpdiff/validation/experimental.py "
          "outputs once the external predictors are wired in.]")
    _results_table(doc,
        "Table 2. Agreement with experimental data on held-out proteins (reduced χ²).",
        ["Protein", "SAXS χ² (no reg)", "SAXS χ² (reg)", "NMR χ² (no reg)", "NMR χ² (reg)"],
        ["[PED ID]", "[PED ID]", "[PED ID]", "Mean"])

    add_heading(doc, "4.4 Ablations and sensitivity", 2)
    p = doc.add_paragraph()
    ph(p, "[RESULT: sweep λ_max to trace the fidelity/diversity curve; sensitivity to "
          "KDE bandwidth; ESM-2 vs no-conditioning control; effect of the annealing "
          "schedule. A sequence-shuffle control (generate with shuffled sequence) "
          "demonstrates the model uses sequence information.]")

    # ---- 5. Discussion ----
    add_heading(doc, "5. Discussion", 1)
    body(doc,
         "We summarise what the regularizer does and does not buy, interpret the "
         "fidelity/diversity trade-off, and discuss limitations: ESM embeddings carry "
         "less structural signal for low-complexity disordered sequences; "
         "torsion-only reconstruction does not model side chains or detect steric "
         "clashes (hence the reported clash rate); and the KDE encodes marginal, not "
         "joint or sequence-specific, angle statistics by design.")
    p = doc.add_paragraph()
    ph(p, "[DISCUSSION: tie conclusions to the numbers; state the regime in which the "
          "regularizer is worthwhile and when its diversity cost is not justified.]")

    # ---- 6. Conclusion ----
    add_heading(doc, "6. Conclusion", 1)
    p = doc.add_paragraph()
    ph(p, "[One paragraph: the central finding about Ramachandran regularization in "
          "torsion-space IDP generation, and the most useful direction it opens.]")

    # ---- Reproducibility ----
    add_heading(doc, "Reproducibility", 1)
    body(doc,
         "Code, configuration files, and the exact run order are provided with this "
         "work (see the accompanying repository README). Data splits are made at the "
         "protein level and serialised so training and evaluation read identical "
         "data; the with- and without-regularization arms differ only in a single "
         "command-line flag (λ_max).")

    # ---- References ----
    add_heading(doc, "References", 1)
    p = doc.add_paragraph()
    ph(p, "[REFERENCES: PED; ESM-2 (Lin et al. 2023); idpGAN (Janson et al. 2023); "
          "torsional diffusion (Jing et al. 2022); DDPM (Ho et al. 2020); cosine "
          "schedule (Nichol & Dhariwal 2021); DiT/AdaLN (Peebles & Xie 2023); SAXS "
          "predictor (CRYSOL / Pepsi-SAXS); NMR predictor (UCBShift / SPARTA+); "
          "statistical coil / flexible-meccano. Format to the target venue.]")

    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "paper_skeleton.docx")
    doc.save(out_path)
    print(f"saved {out_path}")
    return out_path


def _results_table(doc, caption, headers, row_labels):
    cap = doc.add_paragraph()
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        for par in hdr[j].paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(10)
    for label in row_labels:
        cells = table.add_row().cells
        cells[0].text = label
        for j in range(1, len(headers)):
            par = cells[j].paragraphs[0]
            run = par.add_run("[ ]")
            run.font.color.rgb = PLACEHOLDER
    doc.add_paragraph()


if __name__ == "__main__":
    build()
